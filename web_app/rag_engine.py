#!/usr/bin/env python3
"""
rag_engine.py — BGE-M3 + FAISS 知识库引擎

功能：
  - 用 BGE-M3 对 PDF 文档进行 chunk-level 向量化
  - 用 FAISS 存储向量，支持持久化
  - 对话时检索最相关的 chunk 作为 RAG 上下文

存储位置：seismo_rag/
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import time
from datetime import datetime
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple

# ── 路径常量 ──────────────────────────────────────────────────────────────────

KB_DIR      = Path(__file__).parent.parent / "seismo_rag"
INDEX_FILE  = KB_DIR / "faiss_index.bin"
META_FILE   = KB_DIR / "metadata.json"
DOCS_DIR    = KB_DIR / "docs"

for _d in [KB_DIR, DOCS_DIR]:
    _d.mkdir(parents=True, exist_ok=True)


# ── simple_rag 导入助手 ────────────────────────────────────────────────────────
# rag_engine.py 可能被 `from rag_engine import ...`（sys.path 含 web_app/）
# 或 `from web_app.rag_engine import ...`（sys.path 含项目根目录）两种方式导入。
# 统一用这个函数，避免 "from web_app.simple_rag" 在 Flask 下找不到的问题。

def _get_simple_rag():
    """Return the simple_rag singleton, supporting both sys.path layouts."""
    try:
        from simple_rag import get_simple_rag as _f  # web_app/ in sys.path (Flask default)
        return _f()
    except ImportError:
        from web_app.simple_rag import get_simple_rag as _f  # project root in sys.path
        return _f()

# ── 数据模型 ──────────────────────────────────────────────────────────────────

@dataclass
class DocChunk:
    chunk_id:  str          # 唯一 ID = doc_hash + "_" + chunk_index
    doc_id:    str          # 文档哈希
    doc_name:  str          # 原始文件名
    page:      int          # 页码（0-based）
    text:      str          # chunk 文本内容
    char_start: int = 0

@dataclass
class DocMeta:
    doc_id:     str
    doc_name:   str
    file_path:  str         # 副本在 DOCS_DIR 的路径
    n_pages:    int
    n_chunks:   int
    added_at:   str
    size_bytes: int
    proj_folder: str = ""   # 来源项目文件夹（技能/参考文献），空=手动上传
    source_type: str = "upload"  # "upload" | "skill_docs" | "ref_knowledge"


# ── 文本提取与分块 ─────────────────────────────────────────────────────────────

def _extract_pdf_text(pdf_path: str) -> List[Tuple[int, str]]:
    """
    返回 [(page_index, text), ...] 列表。
    优先使用 pdfminer.six，回退到 PyMuPDF。
    """
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTTextContainer
        pages = []
        for i, layout in enumerate(extract_pages(pdf_path)):
            texts = []
            for element in layout:
                if isinstance(element, LTTextContainer):
                    texts.append(element.get_text())
            pages.append((i, " ".join(texts)))
        return pages
    except ImportError:
        pass

    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        pages = [(i, page.get_text()) for i, page in enumerate(doc)]
        doc.close()
        return pages
    except ImportError:
        raise RuntimeError(
            "需要 pdfminer.six 或 PyMuPDF 来解析 PDF：\n"
            "  pip install pdfminer.six\n"
            "  或 pip install pymupdf"
        )


def _extract_docx_text(path: str) -> List[Tuple[int, str]]:
    """
    解析 DOCX，按标题（Heading 样式）分节，返回 [(section_index, text), ...]。
    每个标题开启新节，节内文本合并为一个条目。
    """
    try:
        from docx import Document
        doc = Document(path)
    except ImportError:
        raise RuntimeError("需要 python-docx 解析 DOCX：pip install python-docx")

    sections: List[Tuple[int, str]] = []
    sec_idx = 0
    buf: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            if buf:
                sections.append((sec_idx, "\n".join(buf)))
                sec_idx += 1
                buf = []
            buf.append(text)
        else:
            buf.append(text)
    if buf:
        sections.append((sec_idx, "\n".join(buf)))

    if not sections:
        # 无结构文档：整体作为单页
        full = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [(0, full)] if full else []
    return sections


def _extract_md_text(path: str) -> List[Tuple[int, str]]:
    """
    解析 Markdown，按一级/二级标题（# / ##）分节。
    返回 [(section_index, text), ...]。
    """
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    sections: List[Tuple[int, str]] = []
    sec_idx = 0
    buf: List[str] = []

    for line in text.splitlines():
        # 遇到一级或二级标题时开启新节
        if re.match(r"^#{1,2}\s", line):
            if buf:
                sections.append((sec_idx, "\n".join(buf)))
                sec_idx += 1
                buf = []
        buf.append(line)
    if buf:
        sections.append((sec_idx, "\n".join(buf)))

    return sections if sections else [(0, text)]


def _extract_txt_text(path: str) -> List[Tuple[int, str]]:
    """
    解析纯文本，按三个以上空行分段，每段约 2 000 字符归为一页。
    返回 [(page_index, text), ...]。
    """
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    paragraphs = [p.strip() for p in re.split(r"\n{3,}", text) if p.strip()]
    if not paragraphs:
        return [(0, text.strip())] if text.strip() else []

    pages: List[Tuple[int, str]] = []
    buf: List[str] = []
    buf_len = 0
    page_idx = 0
    PAGE_SIZE = 2000

    for para in paragraphs:
        if buf_len + len(para) > PAGE_SIZE and buf:
            pages.append((page_idx, "\n\n".join(buf)))
            page_idx += 1
            buf = []
            buf_len = 0
        buf.append(para)
        buf_len += len(para)
    if buf:
        pages.append((page_idx, "\n\n".join(buf)))
    return pages


def _extract_rst_text(path: str) -> List[Tuple[int, str]]:
    """
    解析 reStructuredText，按标题分节，返回 [(section_index, text), ...]。
    去除 RST 专有标记（指令、角色、下划线等），保留纯文本内容。
    """
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()
    sections: List[Tuple[int, str]] = []
    buf: List[str] = []
    sec_idx = 0

    # RST 标题：连续相同符号行（=  -  ~  ^  '  "  #  +  *  . ）作为下划线
    _UNDERLINE_RE = re.compile(r'^([=\-~^"\'#\+\*\.]{3,})\s*$')
    _DIRECTIVE_RE = re.compile(r"^\s*\.\.\s+\S")   # .. directive::
    _ROLE_RE      = re.compile(r":\w+:`[^`]*`")    # :role:`text`
    _BACKTICK_RE  = re.compile(r"`+([^`]+)`+")     # `text`
    _FIELD_RE     = re.compile(r"^\s*:\w[^:]*:")    # :param x:

    def _flush():
        nonlocal sec_idx
        clean = []
        skip_block = False
        for ln in buf:
            if _DIRECTIVE_RE.match(ln):
                skip_block = True
                continue
            if skip_block:
                if ln and not ln[0].isspace():
                    skip_block = False
                else:
                    continue
            ln = _ROLE_RE.sub(lambda m: m.group(0).split("`")[1], ln)
            ln = _BACKTICK_RE.sub(r"\1", ln)
            clean.append(ln)
        content = "\n".join(clean).strip()
        if content:
            sections.append((sec_idx, content))
            sec_idx += 1

    i = 0
    while i < len(lines):
        ln = lines[i]
        # 检测标题：当前行非空，下一行是纯下划线且长度匹配
        if (i + 1 < len(lines)
                and ln.strip()
                and _UNDERLINE_RE.match(lines[i + 1])
                and len(lines[i + 1].strip()) >= len(ln.strip())):
            if buf:
                _flush()
                buf = []
            buf.append(ln)   # heading text
            i += 2           # skip underline
            continue
        buf.append(ln)
        i += 1

    if buf:
        _flush()

    # Fallback: treat whole file as one section
    if not sections:
        sections = [(0, text.strip())]
    return sections


def _extract_html_text(path: str) -> List[Tuple[int, str]]:
    """
    解析 HTML 文件，提取可见文本，按 <h1>/<h2> 标题分节。
    返回 [(section_index, text), ...]。
    """
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        SKIP_TAGS = {"script", "style", "nav", "footer", "header",
                     "meta", "link", "noscript", "aside"}
        BLOCK_TAGS = {"h1", "h2", "h3", "h4", "p", "li", "td", "th",
                      "pre", "code", "dt", "dd", "blockquote"}
        HEADING_TAGS = {"h1", "h2"}

        def __init__(self):
            super().__init__()
            self.sections: List[Tuple[int, str]] = []
            self._buf: List[str] = []
            self._skip_depth = 0
            self._cur_tag = ""
            self._sec_idx = 0
            self._cur_lines: List[str] = []

        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            self._cur_tag = tag
            if tag in self.SKIP_TAGS:
                self._skip_depth += 1
            if tag in self.HEADING_TAGS and self._buf:
                content = " ".join(self._buf).strip()
                if content:
                    self.sections.append((self._sec_idx, content))
                    self._sec_idx += 1
                self._buf = []

        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag in self.SKIP_TAGS and self._skip_depth > 0:
                self._skip_depth -= 1
            if tag in self.BLOCK_TAGS:
                self._buf.append("\n")

        def handle_data(self, data):
            if self._skip_depth > 0:
                return
            text = data.strip()
            if text:
                self._buf.append(text)

        def get_sections(self):
            if self._buf:
                content = " ".join(self._buf).strip()
                if content:
                    self.sections.append((self._sec_idx, content))
            return self.sections if self.sections else [(0, "")]

    html_text = Path(path).read_text(encoding="utf-8", errors="replace")
    extractor = _TextExtractor()
    extractor.feed(html_text)
    sections = extractor.get_sections()
    # Filter empty sections
    sections = [(i, s) for i, s in sections if s.strip()]
    return sections if sections else [(0, "")]


def _chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """简单按字符切分，保留语义边界（按句号换行切）。"""
    # 先按段落/句子切
    paragraphs = re.split(r'\n{2,}|(?<=[。！？.!?])\s', text)
    chunks: List[str] = []
    buf = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(buf) + len(para) <= chunk_size:
            buf += (" " if buf else "") + para
        else:
            if buf:
                chunks.append(buf)
            # 如果单段落超长，强制切分
            while len(para) > chunk_size:
                chunks.append(para[:chunk_size])
                para = para[chunk_size - overlap:]
            buf = para
    if buf:
        chunks.append(buf)
    return [c for c in chunks if len(c.strip()) > 20]


# ── 嵌入模型路径读取 ──────────────────────────────────────────────────────────

def _get_embedding_model_path() -> str:
    """
    读取 ~/.seismicx/config.json 中 embedding.model_path 字段。
    未配置时返回默认值 "BAAI/bge-m3"（从 HuggingFace 下载）。
    本地路径示例：/Users/you/open_models/bge-m3
    """
    try:
        cfg_file = Path.home() / ".seismicx" / "config.json"
        if cfg_file.exists():
            cfg = json.loads(cfg_file.read_text(encoding="utf-8"))
            path = cfg.get("embedding", {}).get("model_path", "").strip()
            if path:
                return path
    except Exception:
        pass
    return "BAAI/bge-m3"


# ── 嵌入模型（BGE-M3）─────────────────────────────────────────────────────────

class EmbeddingModel:
    """
    懒加载 BGE-M3 模型。
    优先使用 FlagEmbedding，回退到 sentence-transformers。
    模型路径从 ~/.seismicx/config.json embedding.model_path 读取，
    默认为 "BAAI/bge-m3"（HuggingFace 自动下载）。
    国内用户可先用 ModelScope 下载到本地，再在知识库页面配置本地路径。
    """

    _instance: Optional["EmbeddingModel"] = None

    def __init__(self):
        self._model = None
        self._backend = None
        self.dim = 1024  # BGE-M3 输出维度

    @classmethod
    def get(cls) -> "EmbeddingModel":
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    @classmethod
    def reset(cls):
        """重置单例，下次 get() 时重新加载（用于更新模型路径后强制重载）。"""
        cls._instance = None

    def _load(self):
        if self._model is not None:
            return

        model_path = _get_embedding_model_path()
        flag_err = None
        st_err   = None

        # ── 尝试 FlagEmbedding ────────────────────────────────────────────────
        try:
            from FlagEmbedding import BGEM3FlagModel
            self._model = BGEM3FlagModel(
                model_path,
                use_fp16=True,
                device="cpu",
            )
            self._backend = "flag"
            return
        except ImportError as e:
            flag_err = f"ImportError: {e}"
        except Exception as e:
            flag_err = f"{type(e).__name__}: {e}"

        # ── 尝试 sentence-transformers（标准加载）────────────────────────────
        try:
            from sentence_transformers import SentenceTransformer
            self._model = SentenceTransformer(model_path, device="cpu")
            self._backend = "st"
            return
        except ImportError as e:
            st_err = f"ImportError: {e}"
        except Exception as e:
            st_err = f"{type(e).__name__}: {e}"

        # ── torch 版本过低时的 safetensors 回退 ──────────────────────────────
        # CVE-2025-32434：torch < 2.6 禁止 torch.load，但 safetensors 格式不受限制。
        # 两个库都报 ValueError 且消息含 "CVE-2025-32434" / "torch.load" → 尝试
        # 强制使用 safetensors 加载，避免必须升级 torch。
        _is_torch_cve = lambda e: (
            "CVE-2025-32434" in str(e) or "torch.load" in str(e) or "weights_only" in str(e)
        )
        if (flag_err and _is_torch_cve(flag_err)) or (st_err and _is_torch_cve(st_err)):
            # 方案 A：sentence-transformers + safetensors 模型参数
            try:
                from sentence_transformers import SentenceTransformer
                self._model = SentenceTransformer(
                    model_path,
                    device="cpu",
                    model_kwargs={"use_safetensors": True},
                )
                self._backend = "st-safetensors"
                return
            except Exception:
                pass

            # 方案 B：直接用 transformers AutoModel + safetensors
            try:
                from transformers import AutoTokenizer, AutoModel
                import torch
                tokenizer = AutoTokenizer.from_pretrained(model_path)
                model     = AutoModel.from_pretrained(
                    model_path,
                    use_safetensors=True,
                    torch_dtype=torch.float32,
                )
                model.eval()

                class _TransformersWrapper:
                    """最小包装：模拟 encode() 接口，与 EmbeddingModel.encode() 兼容。"""
                    def __init__(self, tok, mod):
                        self.tokenizer = tok
                        self.model     = mod

                    def encode(self, texts):
                        import torch, numpy as np
                        inputs = self.tokenizer(
                            texts, padding=True, truncation=True,
                            max_length=512, return_tensors="pt"
                        )
                        with torch.no_grad():
                            out = self.model(**inputs)
                        # mean pooling
                        mask = inputs["attention_mask"].unsqueeze(-1).float()
                        vecs = (out.last_hidden_state * mask).sum(1) / mask.sum(1)
                        v = vecs.numpy()
                        # L2 normalize
                        norms = np.linalg.norm(v, axis=1, keepdims=True)
                        return v / np.maximum(norms, 1e-9)

                self._model   = _TransformersWrapper(tokenizer, model)
                self._backend = "transformers-safetensors"
                return
            except Exception as e2:
                pass  # 仍然失败，走到下面的统一报错

        # ── 统一诊断报错 ──────────────────────────────────────────────────────
        import sys
        python = sys.executable
        torch_cve = (flag_err and _is_torch_cve(flag_err)) or \
                    (st_err   and _is_torch_cve(st_err))

        diag_lines = ["未能加载嵌入模型，诊断信息："]
        if flag_err:
            diag_lines.append(f"  FlagEmbedding         → {flag_err}")
        else:
            diag_lines.append("  FlagEmbedding         → 未安装")
        if st_err:
            diag_lines.append(f"  sentence-transformers → {st_err}")
        else:
            diag_lines.append("  sentence-transformers → 未安装")

        if torch_cve:
            diag_lines += [
                "",
                "【原因】torch 版本过低（CVE-2025-32434 要求 torch >= 2.6）",
                "【解决方案】升级 torch：",
                f"  {python} -m pip install 'torch>=2.6'",
                "",
                "升级后无需重启，下次构建自动生效。",
            ]
        elif not flag_err and not st_err:
            diag_lines += [
                "",
                "请安装嵌入模型库：",
                f"  {python} -m pip install FlagEmbedding",
                f"  # 或",
                f"  {python} -m pip install sentence-transformers",
            ]
        else:
            diag_lines += [
                "",
                "如果已安装但仍报错，尝试升级：",
                f"  {python} -m pip install --upgrade FlagEmbedding sentence-transformers torch",
            ]
        raise RuntimeError("\n".join(diag_lines))

    def encode(self, texts: List[str], batch_size: int = 32) -> "np.ndarray":
        import numpy as np
        self._load()

        if self._backend == "flag":
            result = self._model.encode(
                texts,
                batch_size=batch_size,
                max_length=512,
                return_dense=True,
                return_sparse=False,
                return_colbert_vecs=False,
            )
            vecs = result["dense_vecs"]
        elif self._backend == "transformers-safetensors":
            # _TransformersWrapper.encode() batches internally and already L2-normalizes
            vecs = self._model.encode(texts)
        else:
            vecs = self._model.encode(texts, batch_size=batch_size, normalize_embeddings=True)

        vecs = np.array(vecs, dtype="float32")
        # L2 归一化（内积 = cosine）
        norms = np.linalg.norm(vecs, axis=1, keepdims=True)
        norms = np.where(norms == 0, 1, norms)
        return vecs / norms


# ── FAISS 索引封装 ─────────────────────────────────────────────────────────────

class FaissIndex:
    """轻量 FAISS 封装，支持增量添加与持久化。"""

    def __init__(self, dim: int = 1024):
        import faiss
        self.dim = dim
        self.index = faiss.IndexFlatIP(dim)  # 内积（向量已归一化，等价 cosine）
        self._id_map: List[str] = []         # faiss 位置 → chunk_id

    def add(self, vectors: "np.ndarray", chunk_ids: List[str]):
        import numpy as np
        vecs = np.asarray(vectors, dtype="float32")
        self.index.add(vecs)
        self._id_map.extend(chunk_ids)

    def search(self, query_vec: "np.ndarray", top_k: int = 5) -> List[Tuple[str, float]]:
        import numpy as np
        q = np.asarray(query_vec, dtype="float32").reshape(1, -1)
        scores, indices = self.index.search(q, top_k)
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx < 0 or idx >= len(self._id_map):
                continue
            results.append((self._id_map[idx], float(score)))
        return results

    def save(self, index_path: str, idmap_path: str):
        import faiss, json
        faiss.write_index(self.index, index_path)
        with open(idmap_path, "w") as f:
            json.dump(self._id_map, f)

    @classmethod
    def load(cls, index_path: str, idmap_path: str, dim: int = 1024) -> "FaissIndex":
        import faiss, json
        obj = cls(dim)
        obj.index = faiss.read_index(index_path)
        with open(idmap_path) as f:
            obj._id_map = json.load(f)
        return obj

    @property
    def n_vectors(self) -> int:
        return self.index.ntotal


# ── 主知识库类 ────────────────────────────────────────────────────────────────

class KnowledgeBase:
    """
    统一知识库接口。

    Usage
    -----
    kb = KnowledgeBase.get()
    kb.add_pdf("/path/to/paper.pdf", progress_cb=print)
    chunks = kb.retrieve("如何计算b值", top_k=5)
    """

    _instance: Optional["KnowledgeBase"] = None

    def __init__(self):
        self._chunks: dict[str, DocChunk] = {}   # chunk_id → DocChunk
        self._docs:   dict[str, DocMeta]  = {}   # doc_id   → DocMeta
        self._faiss: Optional[FaissIndex] = None
        self._dirty = False
        self._load_state()

    @classmethod
    def get(cls) -> "KnowledgeBase":
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    # ── 状态持久化 ─────────────────────────────────────────────────────────────

    def _idmap_path(self) -> str:
        return str(KB_DIR / "id_map.json")

    def _load_state(self):
        if META_FILE.exists():
            try:
                raw = json.loads(META_FILE.read_text())
                for did, d in raw.get("docs", {}).items():
                    self._docs[did] = DocMeta(**d)
                for cid, c in raw.get("chunks", {}).items():
                    self._chunks[cid] = DocChunk(**c)
            except Exception as e:
                print(f"Error loading metadata: {e}")

        idmap = self._idmap_path()
        if INDEX_FILE.exists() and os.path.exists(idmap):
            try:
                self._faiss = FaissIndex.load(str(INDEX_FILE), idmap)
            except Exception as e:
                print(f"Error loading FAISS index: {e}")
                self._faiss = None

        # 启动时清理：移除 PDF 文件已不存在的孤立文档记录
        self._cleanup_stale_docs()

    def _cleanup_stale_docs(self):
        """删除物理 PDF 已不存在的文档及其 chunk/向量，保持向量库与文件一致。"""
        stale = [did for did, m in self._docs.items()
                 if m.file_path and not Path(m.file_path).exists()]
        if not stale:
            return
        for did in stale:
            self._docs.pop(did, None)
            to_rm = [cid for cid, c in self._chunks.items() if c.doc_id == did]
            for cid in to_rm:
                del self._chunks[cid]
            print(f"[RAG] 清理孤立文档: {did}")
        self._rebuild_index()
        self._save_state()

    def _save_state(self):
        meta = {
            "docs":   {did: asdict(d) for did, d in self._docs.items()},
            "chunks": {cid: asdict(c) for cid, c in self._chunks.items()},
        }
        META_FILE.write_text(json.dumps(meta, ensure_ascii=False, indent=2))
        if self._faiss:
            self._faiss.save(str(INDEX_FILE), self._idmap_path())
        self._dirty = False

    # ── 文档操作 ─────────────────────────────────────────────────────────────

    def add_pdf(
        self,
        pdf_path: str,
        progress_cb=None,
        chunk_size: int = 500,
        overlap: int = 50,
        proj_folder: str = "",
        source_type: str = "upload",
    ) -> DocMeta:
        """
        解析 PDF → 分块 → BGE-M3 编码 → 加入 FAISS。
        返回 DocMeta。
        """
        def _log(msg):
            if progress_cb:
                progress_cb(msg)

        pdf_path = str(pdf_path)
        doc_name = Path(pdf_path).name

        # 计算文档 hash（含路径特征，防止不同子文件夹同名同内容文件冲突）
        _path_sig = f"{proj_folder}/{doc_name}".encode()
        h = hashlib.md5(_path_sig + b"|" + Path(pdf_path).read_bytes()).hexdigest()[:12]
        doc_id = h

        if doc_id in self._docs:
            _log(f"⚠ 文档已存在于知识库：{doc_name}，跳过")
            return self._docs[doc_id]

        _log(f"📄 解析 PDF：{doc_name}")
        pages = _extract_pdf_text(pdf_path)
        n_pages = len(pages)
        _log(f"   共 {n_pages} 页")

        # 分块
        all_chunks: List[DocChunk] = []
        for page_idx, page_text in pages:
            sub_chunks = _chunk_text(page_text, chunk_size, overlap)
            for i, chunk_text in enumerate(sub_chunks):
                cid = f"{doc_id}_{page_idx}_{i}"
                all_chunks.append(DocChunk(
                    chunk_id=cid,
                    doc_id=doc_id,
                    doc_name=doc_name,
                    page=page_idx,
                    text=chunk_text,
                ))
        _log(f"   切分为 {len(all_chunks)} 个 chunk")

        if not all_chunks:
            raise ValueError("PDF 中未提取到有效文本，可能是扫描版 PDF（需 OCR）")

        # 编码
        _log("🔢 BGE-M3 向量化中…")
        try:
            model = EmbeddingModel.get()
            texts  = [c.text for c in all_chunks]
            vecs   = model.encode(texts)
            _log(f"   向量维度: {vecs.shape}")

            # 加入 FAISS
            if self._faiss is None:
                self._faiss = FaissIndex(dim=vecs.shape[1])
            chunk_ids = [c.chunk_id for c in all_chunks]
            self._faiss.add(vecs, chunk_ids)

            # 注册 chunk
            for chunk in all_chunks:
                self._chunks[chunk.chunk_id] = chunk
        except RuntimeError as e:
            _log(f"⚠ 嵌入模型加载失败: {e}")
            _log("⚠ 正在使用简化版RAG实现...")
            
            # 使用简化版实现
            simple_rag = _get_simple_rag()
            
            # 添加文档到简化版RAG
            doc_metadata = {
                "doc_id": doc_id,
                "doc_name": doc_name,
                "file_path": "",  # 简化版不需要存储文件
                "n_pages": n_pages,
                "n_chunks": len(all_chunks),
                "added_at": datetime.now().isoformat(timespec="seconds"),
                "size_bytes": Path(pdf_path).stat().st_size,
            }
            simple_rag.add_document(all_chunks, doc_metadata)
            
            _log(f"✅ 已使用简化版RAG添加文档：{doc_name}（{len(all_chunks)} chunks）")
            meta = DocMeta(
                doc_id=doc_id,
                doc_name=doc_name,
                file_path="",  # 简化版不需要存储文件
                n_pages=n_pages,
                n_chunks=len(all_chunks),
                added_at=datetime.now().isoformat(timespec="seconds"),
                size_bytes=Path(pdf_path).stat().st_size,
                proj_folder=proj_folder,
                source_type=source_type,
            )
            self._docs[doc_id] = meta
            # 确保元数据持久化
            self._save_state()
            return meta

        # 复制 PDF 到知识库目录
        dest = DOCS_DIR / f"{doc_id}_{doc_name}"
        shutil.copy2(pdf_path, dest)

        # 注册文档
        meta = DocMeta(
            doc_id=doc_id,
            doc_name=doc_name,
            file_path=str(dest),
            n_pages=n_pages,
            n_chunks=len(all_chunks),
            added_at=datetime.now().isoformat(timespec="seconds"),
            size_bytes=Path(pdf_path).stat().st_size,
            proj_folder=proj_folder,
            source_type=source_type,
        )
        self._docs[doc_id] = meta
        self._save_state()
        _log(f"✅ 已加入知识库：{doc_name}（{len(all_chunks)} chunks）")
        return meta

    # ── 通用多格式文档入库 ────────────────────────────────────────────────────

    def add_document(
        self,
        path: str,
        progress_cb=None,
        chunk_size: int = 500,
        overlap: int = 50,
        proj_folder: str = "",
        source_type: str = "upload",
    ) -> "DocMeta":
        """
        将任意受支持格式的文档加入知识库。

        支持格式
        --------
        .pdf   — PDF（调用 add_pdf，与直接上传行为完全一致）
        .docx  — Word 文档（按标题分节）
        .md    — Markdown（按一级/二级标题分节）
        .txt   — 纯文本（按空行分页）

        Parameters
        ----------
        path : str   文件绝对路径
        progress_cb  进度回调 callable(msg: str)
        """
        ext = Path(path).suffix.lower()
        if ext == ".pdf":
            return self.add_pdf(path, progress_cb, chunk_size, overlap,
                                proj_folder=proj_folder, source_type=source_type)
        elif ext == ".docx":
            pages = _extract_docx_text(path)
        elif ext == ".md":
            pages = _extract_md_text(path)
        elif ext in (".txt", ".text"):
            pages = _extract_txt_text(path)
        elif ext == ".rst":
            pages = _extract_rst_text(path)
        elif ext in (".html", ".htm"):
            pages = _extract_html_text(path)
        else:
            raise ValueError(
                f"不支持的文件类型：{ext}\n"
                "支持：.pdf  .docx  .txt  .md  .rst  .html"
            )
        return self._add_pages(path, pages, progress_cb, chunk_size, overlap,
                               proj_folder=proj_folder, source_type=source_type)

    def _add_pages(
        self,
        file_path: str,
        pages: List[Tuple[int, str]],
        progress_cb=None,
        chunk_size: int = 500,
        overlap: int = 50,
        proj_folder: str = "",
        source_type: str = "upload",
    ) -> "DocMeta":
        """
        通用索引流水线：将预提取的 (section_idx, text) 列表向量化并存入知识库。
        被 add_document() 调用，供非 PDF 格式使用。
        """
        def _log(msg: str):
            if progress_cb:
                progress_cb(msg)

        file_path = str(file_path)
        doc_name  = Path(file_path).name
        # 含路径特征，防止不同子文件夹同名同内容文件冲突
        _path_sig = f"{proj_folder}/{doc_name}".encode()
        h         = hashlib.md5(_path_sig + b"|" + Path(file_path).read_bytes()).hexdigest()[:12]
        doc_id    = h

        if doc_id in self._docs:
            _log(f"⚠ 文档已存在于知识库：{doc_name}，跳过")
            return self._docs[doc_id]

        _log(f"📄 解析文档：{doc_name}")
        n_pages = len(pages)
        _log(f"   共 {n_pages} 节/页")

        # 分块
        all_chunks: List[DocChunk] = []
        for page_idx, page_text in pages:
            for i, ct in enumerate(_chunk_text(page_text, chunk_size, overlap)):
                cid = f"{doc_id}_{page_idx}_{i}"
                all_chunks.append(DocChunk(
                    chunk_id=cid, doc_id=doc_id, doc_name=doc_name,
                    page=page_idx, text=ct,
                ))
        _log(f"   切分为 {len(all_chunks)} 个 chunk")

        if not all_chunks:
            raise ValueError("文档中未提取到有效文本（可能为空文件或编码问题）")

        # 向量化并加入 FAISS
        _log("🔢 BGE-M3 向量化中…")
        try:
            model     = EmbeddingModel.get()
            texts     = [c.text for c in all_chunks]
            vecs      = model.encode(texts)
            _log(f"   向量维度: {vecs.shape}")
            if self._faiss is None:
                self._faiss = FaissIndex(dim=vecs.shape[1])
            self._faiss.add(vecs, [c.chunk_id for c in all_chunks])
            for chunk in all_chunks:
                self._chunks[chunk.chunk_id] = chunk
        except RuntimeError as e:
            _log(f"⚠ 嵌入模型加载失败: {e}")
            _log("⚠ 正在使用简化版 RAG…")
            simple_rag = _get_simple_rag()
            simple_rag.add_document(all_chunks, {
                "doc_id": doc_id, "doc_name": doc_name,
                "file_path": "", "n_pages": n_pages,
                "n_chunks": len(all_chunks),
                "added_at": datetime.now().isoformat(timespec="seconds"),
                "size_bytes": Path(file_path).stat().st_size,
                "proj_folder": proj_folder,
                "source_type": source_type,
            })
            meta = DocMeta(
                doc_id=doc_id, doc_name=doc_name, file_path="",
                n_pages=n_pages, n_chunks=len(all_chunks),
                added_at=datetime.now().isoformat(timespec="seconds"),
                size_bytes=Path(file_path).stat().st_size,
                proj_folder=proj_folder,
                source_type=source_type,
            )
            self._docs[doc_id] = meta
            self._save_state()
            _log(f"✅ 已加入知识库（简化版）：{doc_name}（{len(all_chunks)} chunks）")
            return meta

        # 保存文件副本
        dest = DOCS_DIR / f"{doc_id}_{doc_name}"
        shutil.copy2(file_path, dest)

        meta = DocMeta(
            doc_id=doc_id, doc_name=doc_name, file_path=str(dest),
            n_pages=n_pages, n_chunks=len(all_chunks),
            added_at=datetime.now().isoformat(timespec="seconds"),
            size_bytes=Path(file_path).stat().st_size,
            proj_folder=proj_folder,
            source_type=source_type,
        )
        self._docs[doc_id] = meta
        self._save_state()
        _log(f"✅ 已加入知识库：{doc_name}（{len(all_chunks)} chunks）")
        return meta

    def delete_doc(self, doc_id: str) -> bool:
        """
        删除文档及其所有 chunk。
        注意：FAISS IndexFlatIP 不支持按 ID 删除，需要重建索引。
        """
        if doc_id in self._docs:
            # 标准FAISS实现
            meta = self._docs.pop(doc_id)
            # 移除 chunk
            to_remove = [cid for cid, c in self._chunks.items() if c.doc_id == doc_id]
            for cid in to_remove:
                del self._chunks[cid]

            # 删除文件副本
            try:
                Path(meta.file_path).unlink(missing_ok=True)
            except Exception:
                pass

            # 重建 FAISS 索引（只保留剩余 chunk 的向量）
            self._rebuild_index()
            self._save_state()
            return True
        else:
            # 检查简化版RAG系统中是否存在该文档
            try:
                simple_rag = _get_simple_rag()
                if doc_id in simple_rag._docs:
                    # 从简化版RAG中删除文档
                    result = simple_rag.delete_document(doc_id)
                    return result
                else:
                    return False
            except Exception as e:
                print(f"Error deleting document from simple RAG: {e}")
                return False

    def _rebuild_index(self):
        """从剩余 chunks 重建 FAISS 索引。"""
        if not self._chunks:
            self._faiss = None
            return

        model  = EmbeddingModel.get()
        cids   = list(self._chunks.keys())
        texts  = [self._chunks[cid].text for cid in cids]
        vecs   = model.encode(texts)
        dim    = vecs.shape[1]
        self._faiss = FaissIndex(dim=dim)
        self._faiss.add(vecs, cids)

    def clear(self):
        """清空整个知识库（FAISS + SimpleRAG 两个存储后端同时清空）。"""
        self._chunks.clear()
        self._docs.clear()
        self._faiss = None
        for f in DOCS_DIR.iterdir():
            f.unlink(missing_ok=True)
        INDEX_FILE.unlink(missing_ok=True)
        META_FILE.unlink(missing_ok=True)
        Path(self._idmap_path()).unlink(missing_ok=True)
        # 同时清空 SimpleRAG（TF-IDF 回退后端），否则 list_docs/status 会从那里
        # 读回"已清空"的文档（因为此时 _faiss is None，触发回退逻辑）
        try:
            _get_simple_rag().clear()
        except Exception:
            pass

    # ── 检索 ─────────────────────────────────────────────────────────────────

    def retrieve(
        self,
        query: str,
        top_k: int = 5,
        score_threshold: float = 0.3,
    ) -> List[Tuple[DocChunk, float]]:
        """
        检索最相关的 chunk。
        返回 [(DocChunk, score), ...]，按 score 降序。
        """
        # 如果 FAISS 不可用，使用简化版 RAG
        if not self._faiss or self._faiss.n_vectors == 0:
            try:
                simple_rag = _get_simple_rag()
                results = simple_rag.retrieve(query, top_k)

                # TF-IDF 分数远低于向量余弦相似度，使用更宽松的阈值 0.05
                tfidf_threshold = min(score_threshold, 0.05)
                formatted_results = []
                for text, score, metadata in results:
                    if score >= tfidf_threshold:
                        chunk = DocChunk(
                            chunk_id=metadata.get("chunk_id", "virtual"),
                            doc_id=metadata.get("doc_id", "virtual"),
                            doc_name=metadata.get("doc_name", "Virtual Doc"),
                            page=metadata.get("page", 0),
                            text=text
                        )
                        formatted_results.append((chunk, score))

                return formatted_results
            except Exception:
                # 如果简化版也失败了，返回空结果
                return []

        try:
            model = EmbeddingModel.get()
            qvec  = model.encode([query])
            hits  = self._faiss.search(qvec, top_k=top_k)

            results = []
            for cid, score in hits:
                if score < score_threshold:
                    continue
                if cid in self._chunks:
                    results.append((self._chunks[cid], score))
            return sorted(results, key=lambda x: x[1], reverse=True)
        except RuntimeError:
            # 如果嵌入模型不可用，使用简化版 RAG
            try:
                simple_rag = _get_simple_rag()
                results = simple_rag.retrieve(query, top_k)

                tfidf_threshold = min(score_threshold, 0.05)
                formatted_results = []
                for text, score, metadata in results:
                    if score >= tfidf_threshold:
                        chunk = DocChunk(
                            chunk_id=metadata.get("chunk_id", "virtual"),
                            doc_id=metadata.get("doc_id", "virtual"),
                            doc_name=metadata.get("doc_name", "Virtual Doc"),
                            page=metadata.get("page", 0),
                            text=text
                        )
                        formatted_results.append((chunk, score))

                return formatted_results
            except Exception:
                return []

    def retrieve_relevant_docs(
        self,
        query: str,
        top_k: int = 8,
        score_threshold: float = 0.5,
    ) -> List[dict]:
        """
        直接检索高度相关文献，返回结构化结果列表。

        每条结果包含：
          - doc_name: 文献文件名
          - page:     页码（1-based）
          - score:    相关度得分（0~1）
          - text:     命中文本段落
          - chunk_id: chunk 唯一 ID
          - doc_id:   文档 ID

        Parameters
        ----------
        query          : 检索查询
        top_k          : 最多返回条数（超过阈值的才保留）
        score_threshold: 相关度阈值，默认 0.5（仅返回高度相关结果）
        """
        hits = self.retrieve(query, top_k=top_k, score_threshold=score_threshold)
        results = []
        for chunk, score in hits:
            results.append({
                "doc_name":  chunk.doc_name,
                "page":      chunk.page + 1,   # 转为 1-based
                "score":     round(score, 4),
                "text":      chunk.text,
                "chunk_id":  chunk.chunk_id,
                "doc_id":    chunk.doc_id,
            })
        return results

    def build_rag_context(
        self,
        query: str,
        top_k: int = 5,
        max_chars: int = 3000,
        score_threshold: float = 0.5,
    ) -> str:
        """
        检索 + 格式化为 LLM 系统提示可用的上下文段落。
        """
        # 如果 FAISS 不可用，使用简化版 RAG
        if not self._faiss or self._faiss.n_vectors == 0:
            try:
                simple_rag = _get_simple_rag()
                return simple_rag.build_context(query, top_k, max_chars)
            except Exception:
                return ""

        hits = self.retrieve(query, top_k=top_k, score_threshold=score_threshold)
        if not hits:
            try:
                # 尝试使用简化版 RAG
                simple_rag = _get_simple_rag()
                return simple_rag.build_context(query, top_k, max_chars)
            except Exception:
                return ""

        lines = ["以下是从知识库中检索到的相关内容，请参考这些内容回答用户问题：\n"]
        total = 0
        for chunk, score in hits:
            entry = (
                f"【来源：{chunk.doc_name}，第 {chunk.page + 1} 页，"
                f"相关度 {score:.2f}】\n{chunk.text}\n"
            )
            if total + len(entry) > max_chars:
                break
            lines.append(entry)
            total += len(entry)

        return "\n".join(lines)

    # ── 状态查询 ──────────────────────────────────────────────────────────────

    def list_docs(self) -> List[DocMeta]:
        # 获取常规文档
        docs = {**self._docs}
        
        # 如果FAISS索引为空，尝试从简化版RAG获取文档
        if not self._faiss or self._faiss.n_vectors == 0:
            try:
                simple_rag = _get_simple_rag()
                # 将简化版文档合并进来
                for doc_id, doc_meta in simple_rag._docs.items():
                    if doc_id not in docs:
                        docs[doc_id] = doc_meta
            except Exception:
                pass  # 如果简化版RAG不可用，只返回常规文档
        
        return sorted(docs.values(), key=lambda d: d.added_at, reverse=True)

    @property
    def n_docs(self) -> int:
        count = len(self._docs)
        
        # 如果FAISS索引为空，尝试从简化版RAG获取文档数
        if not self._faiss or self._faiss.n_vectors == 0:
            try:
                simple_rag = _get_simple_rag()
                count = max(count, len(simple_rag._docs))
            except Exception:
                pass
        
        return count

    @property
    def n_chunks(self) -> int:
        count = len(self._chunks)
        
        # 如果FAISS索引为空，尝试从简化版RAG获取块数
        if not self._faiss or self._faiss.n_vectors == 0:
            try:
                simple_rag = _get_simple_rag()
                # 简化版RAG的块数可以通过其向量数据库获取
                count = max(count, simple_rag.db.count_items())
            except Exception:
                pass
        
        return count

    @property
    def is_empty(self) -> bool:
        return self.n_chunks == 0

    def status(self) -> dict:
        n_vectors = self._faiss.n_vectors if self._faiss else 0
        
        # 如果FAISS索引为空，尝试从简化版RAG获取向量数
        if n_vectors == 0:
            try:
                simple_rag = _get_simple_rag()
                n_vectors = simple_rag.db.count_items()
            except Exception:
                pass
        
        return {
            "n_docs":   self.n_docs,
            "n_chunks": self.n_chunks,
            "n_vectors": n_vectors,
            "kb_dir":   str(KB_DIR),
        }


# ── 单例访问 ──────────────────────────────────────────────────────────────────

def get_knowledge_base() -> KnowledgeBase:
    return KnowledgeBase.get()


if __name__ == "__main__":
    import sys
    kb = get_knowledge_base()
    print("知识库状态:", kb.status())

    if len(sys.argv) > 1:
        pdf = sys.argv[1]
        kb.add_pdf(pdf, progress_cb=print)
        print("\n检索测试:")
        hits = kb.retrieve("地震波形处理", top_k=3)
        for chunk, score in hits:
            print(f"  [{score:.3f}] {chunk.doc_name} p{chunk.page}: {chunk.text[:80]}...")