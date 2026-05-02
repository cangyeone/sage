"""
rag_extractors.py — Text extraction and chunking helpers for the RAG pipeline.

Each extractor returns a list of (section_index, text) tuples that map naturally
to pages (PDF) or logical sections (Markdown headings, DOCX headings, etc.).

Public API
----------
extract_text(path)                → List[Tuple[int, str]]
chunk_text(text, size, overlap)   → List[str]

Supported formats: .pdf  .docx  .md  .txt  .rst  .html  .htm
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Tuple


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

_EXTRACTORS = {}  # populated at the bottom of this file


def extract_text(path: str) -> List[Tuple[int, str]]:
    """
    Dispatch to the appropriate extractor based on file extension.

    Returns
    -------
    List of (section_index, text) pairs.
    section_index is 0-based (page for PDFs, logical section for others).

    Raises
    ------
    ValueError  — unsupported file extension
    RuntimeError — missing optional dependency
    """
    ext = Path(path).suffix.lower()
    fn = _EXTRACTORS.get(ext)
    if fn is None:
        supported = "  ".join(sorted(_EXTRACTORS))
        raise ValueError(
            f"Unsupported file type: {ext!r}\n"
            f"Supported: {supported}"
        )
    return fn(path)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _extract_pdf(path: str) -> List[Tuple[int, str]]:
    """
    Extract text page-by-page.
    Tries pdfminer.six first; falls back to PyMuPDF (fitz).
    """
    # --- pdfminer.six -------------------------------------------------------
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTTextContainer

        pages = []
        for i, layout in enumerate(extract_pages(path)):
            texts = [el.get_text() for el in layout if isinstance(el, LTTextContainer)]
            pages.append((i, " ".join(texts)))
        return pages
    except ImportError:
        pass

    # --- PyMuPDF (fitz) -----------------------------------------------------
    try:
        import fitz  # type: ignore
        doc = fitz.open(path)
        pages = [(i, page.get_text()) for i, page in enumerate(doc)]
        doc.close()
        return pages
    except ImportError:
        pass

    raise RuntimeError(
        "PDF parsing requires pdfminer.six or PyMuPDF:\n"
        "  pip install pdfminer.six\n"
        "  # or\n"
        "  pip install pymupdf"
    )


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _extract_docx(path: str) -> List[Tuple[int, str]]:
    """
    Split DOCX by Heading styles; each heading starts a new section.
    Falls back to whole-document text if no headings are found.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise RuntimeError(
            "DOCX parsing requires python-docx:\n"
            "  pip install python-docx"
        )

    doc = Document(path)
    sections: List[Tuple[int, str]] = []
    idx = 0
    buf: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            if buf:
                sections.append((idx, "\n".join(buf)))
                idx += 1
                buf = []
        buf.append(text)

    if buf:
        sections.append((idx, "\n".join(buf)))

    if not sections:
        full = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [(0, full)] if full else []

    return sections


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

def _extract_md(path: str) -> List[Tuple[int, str]]:
    """Split Markdown on H1/H2 headers (# / ##)."""
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    sections: List[Tuple[int, str]] = []
    idx = 0
    buf: List[str] = []

    for line in text.splitlines():
        if re.match(r"^#{1,2}\s", line):
            if buf:
                sections.append((idx, "\n".join(buf)))
                idx += 1
                buf = []
        buf.append(line)

    if buf:
        sections.append((idx, "\n".join(buf)))

    return sections if sections else [(0, text)]


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------

def _extract_txt(path: str) -> List[Tuple[int, str]]:
    """
    Split plain text on triple blank lines; group paragraphs into ~2 000-char pages.
    """
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    paragraphs = [p.strip() for p in re.split(r"\n{3,}", text) if p.strip()]
    if not paragraphs:
        return [(0, text.strip())] if text.strip() else []

    PAGE_SIZE = 2000
    pages: List[Tuple[int, str]] = []
    buf: List[str] = []
    buf_len = 0
    page_idx = 0

    for para in paragraphs:
        if buf_len + len(para) > PAGE_SIZE and buf:
            pages.append((page_idx, "\n\n".join(buf)))
            page_idx += 1
            buf = []
            buf_len = 0
        buf.append(para)
        buf_len += len(para)

    if buf:
        pages.append((page_idx, "\n\n".join(buf)))

    return pages


# ---------------------------------------------------------------------------
# reStructuredText
# ---------------------------------------------------------------------------

_RST_UNDERLINE = re.compile(r'^([=\-~^"\'#\+\*\.]{3,})\s*$')
_RST_DIRECTIVE = re.compile(r"^\s*\.\.\s+\S")
_RST_ROLE      = re.compile(r":\w+:`([^`]*)`")
_RST_BACKTICK  = re.compile(r"`+([^`]+)`+")


def _extract_rst(path: str) -> List[Tuple[int, str]]:
    """
    Parse reStructuredText: split on underlined headings, strip RST markup.
    """
    text  = Path(path).read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()
    sections: List[Tuple[int, str]] = []
    buf:  List[str] = []
    idx   = 0

    def _flush():
        nonlocal idx
        cleaned, skip = [], False
        for ln in buf:
            if _RST_DIRECTIVE.match(ln):
                skip = True
                continue
            if skip:
                if ln and not ln[0].isspace():
                    skip = False
                else:
                    continue
            ln = _RST_ROLE.sub(r"\1", ln)
            ln = _RST_BACKTICK.sub(r"\1", ln)
            cleaned.append(ln)
        content = "\n".join(cleaned).strip()
        if content:
            sections.append((idx, content))
            idx += 1

    i = 0
    while i < len(lines):
        ln = lines[i]
        if (
            i + 1 < len(lines)
            and ln.strip()
            and _RST_UNDERLINE.match(lines[i + 1])
            and len(lines[i + 1].strip()) >= len(ln.strip())
        ):
            if buf:
                _flush()
                buf = []
            buf.append(ln)
            i += 2
            continue
        buf.append(ln)
        i += 1

    if buf:
        _flush()

    return sections if sections else [(0, text.strip())]


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

def _extract_html(path: str) -> List[Tuple[int, str]]:
    """Extract visible text from HTML; split on <h1>/<h2> boundaries."""
    from html.parser import HTMLParser

    class _Extractor(HTMLParser):
        _SKIP  = {"script", "style", "nav", "footer", "header",
                  "meta", "link", "noscript", "aside"}
        _BLOCK = {"h1", "h2", "h3", "h4", "p", "li", "td", "th",
                  "pre", "code", "dt", "dd", "blockquote"}
        _HEAD  = {"h1", "h2"}

        def __init__(self):
            super().__init__()
            self._sections: List[Tuple[int, str]] = []
            self._buf: List[str] = []
            self._skip_depth = 0
            self._sec_idx = 0

        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            if tag in self._SKIP:
                self._skip_depth += 1
            if tag in self._HEAD and self._buf:
                self._flush()

        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag in self._SKIP and self._skip_depth > 0:
                self._skip_depth -= 1
            if tag in self._BLOCK:
                self._buf.append("\n")

        def handle_data(self, data):
            if self._skip_depth == 0:
                t = data.strip()
                if t:
                    self._buf.append(t)

        def _flush(self):
            content = " ".join(self._buf).strip()
            if content:
                self._sections.append((self._sec_idx, content))
                self._sec_idx += 1
            self._buf = []

        def result(self):
            self._flush()
            return [(i, s) for i, s in self._sections if s.strip()]

    ext = _Extractor()
    ext.feed(Path(path).read_text(encoding="utf-8", errors="replace"))
    sections = ext.result()
    return sections if sections else [(0, "")]


# ---------------------------------------------------------------------------
# Register extractors
# ---------------------------------------------------------------------------

_EXTRACTORS = {
    ".pdf":  _extract_pdf,
    ".docx": _extract_docx,
    ".md":   _extract_md,
    ".txt":  _extract_txt,
    ".text": _extract_txt,
    ".rst":  _extract_rst,
    ".html": _extract_html,
    ".htm":  _extract_html,
}


# ---------------------------------------------------------------------------
# Chunker
# ---------------------------------------------------------------------------

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """
    Split text into overlapping chunks of at most `chunk_size` characters.

    Strategy: split on paragraph/sentence boundaries first, then hard-cut
    any single paragraph that still exceeds `chunk_size`.

    Parameters
    ----------
    text       : source text
    chunk_size : target maximum characters per chunk
    overlap    : characters carried forward from the previous chunk on hard cuts

    Returns
    -------
    List of non-empty strings, each at least 20 characters long.
    """
    paragraphs = re.split(r'\n{2,}|(?<=[。！？.!?])\s', text)
    chunks: List[str] = []
    buf = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(buf) + len(para) <= chunk_size:
            buf = (buf + " " + para).strip() if buf else para
        else:
            if buf:
                chunks.append(buf)
            while len(para) > chunk_size:
                chunks.append(para[:chunk_size])
                para = para[chunk_size - overlap:]
            buf = para

    if buf:
        chunks.append(buf)

    return [c for c in chunks if len(c.strip()) >= 20]
